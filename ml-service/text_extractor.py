"""
Text Extractor Module

Extracts text content from PDF and DOCX documents for EIR analysis.
Supports chunking for large documents that exceed LLM context windows.
"""

import os
import io
import logging
from typing import List, Tuple, Optional
from pathlib import Path

# PDF extraction
import pdfplumber

# DOCX extraction
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


class TextExtractor:
    """
    Extracts text from PDF and DOCX files with support for:
    - Multi-page PDFs
    - Tables in PDFs (converted to text)
    - DOCX paragraphs and tables
    - Text chunking for large documents
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}

    def __init__(self, max_chunk_tokens: int = 4000):
        """
        Initialize the text extractor.

        Args:
            max_chunk_tokens: Maximum tokens per chunk (approximate, using 4 chars per token)
        """
        self.max_chunk_tokens = max_chunk_tokens
        self.chars_per_token = 4  # Rough estimate

    def extract(self, file_path: str = None, file_bytes: bytes = None,
                filename: str = None) -> Tuple[str, dict]:
        """
        Extract text from a file (path or bytes).

        Args:
            file_path: Path to the file
            file_bytes: Raw file bytes
            filename: Original filename (required if using file_bytes)

        Returns:
            Tuple of (extracted_text, metadata_dict)
        """
        if file_path:
            ext = Path(file_path).suffix.lower()
            with open(file_path, 'rb') as f:
                file_bytes = f.read()
        elif file_bytes and filename:
            ext = Path(filename).suffix.lower()
        else:
            raise ValueError("Either file_path or (file_bytes + filename) required")

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")

        if ext == '.pdf':
            return self.extract_from_pdf(file_bytes)
        elif ext in {'.docx', '.doc'}:
            return self.extract_from_docx(file_bytes)

    def extract_from_pdf(self, file_bytes: bytes) -> Tuple[str, dict]:
        """
        Extract text from a PDF file.

        Args:
            file_bytes: PDF file content as bytes

        Returns:
            Tuple of (extracted_text, metadata)
        """
        text_parts = []
        metadata = {
            'pages': 0,
            'tables_found': 0,
            'extraction_method': 'pdfplumber'
        }

        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                metadata['pages'] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = []

                    # Extract main text
                    text = page.extract_text()
                    if text:
                        page_text.append(text)

                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        metadata['tables_found'] += len(tables)
                        for table in tables:
                            table_text = self._table_to_text(table)
                            if table_text:
                                page_text.append(f"\n[Table]\n{table_text}\n[/Table]")

                    if page_text:
                        text_parts.append(f"\n--- Page {page_num} ---\n")
                        text_parts.extend(page_text)

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise RuntimeError(f"Failed to extract text from PDF: {e}")

        full_text = '\n'.join(text_parts)
        metadata['word_count'] = len(full_text.split())
        metadata['char_count'] = len(full_text)

        return full_text, metadata

    def extract_from_docx(self, file_bytes: bytes) -> Tuple[str, dict]:
        """
        Extract text from a DOCX file.

        Args:
            file_bytes: DOCX file content as bytes

        Returns:
            Tuple of (extracted_text, metadata)
        """
        text_parts = []
        metadata = {
            'paragraphs': 0,
            'tables_found': 0,
            'extraction_method': 'python-docx'
        }

        try:
            doc = DocxDocument(io.BytesIO(file_bytes))

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    metadata['paragraphs'] += 1

            # Extract tables
            for table in doc.tables:
                metadata['tables_found'] += 1
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                table_text = self._table_to_text(table_data)
                if table_text:
                    text_parts.append(f"\n[Table]\n{table_text}\n[/Table]")

        except PackageNotFoundError:
            raise RuntimeError("Invalid DOCX file format")
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise RuntimeError(f"Failed to extract text from DOCX: {e}")

        full_text = '\n\n'.join(text_parts)
        metadata['word_count'] = len(full_text.split())
        metadata['char_count'] = len(full_text)
        metadata['pages'] = max(1, metadata['paragraphs'] // 25)  # Rough estimate

        return full_text, metadata

    def _table_to_text(self, table: List[List[str]]) -> str:
        """
        Convert a table (list of rows) to formatted text.

        Args:
            table: 2D list of cell values

        Returns:
            Formatted table as text
        """
        if not table:
            return ""

        # Filter empty rows
        table = [row for row in table if any(cell and cell.strip() for cell in row)]
        if not table:
            return ""

        # Format as simple text table
        lines = []
        for row in table:
            cells = [str(cell).strip() if cell else '' for cell in row]
            lines.append(' | '.join(cells))

        return '\n'.join(lines)

    def chunk_text(self, text: str, overlap_tokens: int = 200) -> List[str]:
        """
        Split text into chunks that fit within the token limit.

        Args:
            text: Full text to chunk
            overlap_tokens: Number of tokens to overlap between chunks

        Returns:
            List of text chunks
        """
        max_chars = self.max_chunk_tokens * self.chars_per_token
        overlap_chars = overlap_tokens * self.chars_per_token

        if len(text) <= max_chars:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + max_chars

            # Try to break at paragraph boundary
            if end < len(text):
                # Look for paragraph break
                para_break = text.rfind('\n\n', start, end)
                if para_break > start + (max_chars // 2):
                    end = para_break + 2
                else:
                    # Look for sentence break
                    sentence_break = max(
                        text.rfind('. ', start, end),
                        text.rfind('.\n', start, end),
                        text.rfind('? ', start, end),
                        text.rfind('! ', start, end)
                    )
                    if sentence_break > start + (max_chars // 2):
                        end = sentence_break + 2

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - overlap_chars
            if start >= len(text):
                break

        return chunks

    def estimate_tokens(self, text: str) -> int:
        """
        Estimate the number of tokens in text.

        Args:
            text: Text to estimate

        Returns:
            Estimated token count
        """
        return len(text) // self.chars_per_token


def create_extractor(max_chunk_tokens: int = 4000) -> TextExtractor:
    """Factory function to create a TextExtractor instance."""
    return TextExtractor(max_chunk_tokens=max_chunk_tokens)


# Module-level singleton for convenience
_extractor = None


def get_extractor() -> TextExtractor:
    """Get or create the singleton TextExtractor instance."""
    global _extractor
    if _extractor is None:
        _extractor = TextExtractor()
    return _extractor


if __name__ == "__main__":
    # Test extraction
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        extractor = TextExtractor()
        text, metadata = extractor.extract(file_path=file_path)

        print(f"Extracted {metadata.get('pages', 'N/A')} pages")
        print(f"Word count: {metadata.get('word_count', 'N/A')}")
        print(f"Tables found: {metadata.get('tables_found', 0)}")
        print("\n--- First 1000 characters ---\n")
        print(text[:1000])
    else:
        print("Usage: python text_extractor.py <file_path>")
